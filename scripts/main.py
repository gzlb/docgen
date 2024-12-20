from pathlib import Path
import pandas as pd
from docx import Document
import os

def read_parameters(file_path: Path) -> dict[str, str]:
    """Read parameters from an Excel file."""
    df = pd.read_excel(file_path)
    return dict(zip(df['Parameter Name'], df['Value']))

def update_paragraphs(paragraphs, parameters: dict[str, str]) -> None:
    """Replace placeholders in paragraphs."""
    for paragraph in paragraphs:
        for key, value in parameters.items():
            if f"{{{key}}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"{{{key}}}", str(value))

def update_tables(tables, parameters: dict[str, str]) -> None:
    """Replace placeholders in tables."""
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                update_paragraphs(cell.paragraphs, parameters)

def update_headers_footers(doc, parameters: dict[str, str]) -> None:
    """Replace placeholders in headers and footers."""
    for section in doc.sections:
        update_paragraphs(section.header.paragraphs, parameters)
        update_paragraphs(section.footer.paragraphs, parameters)

def update_inline_shapes(doc, parameters: dict[str, str]) -> None:
    """Replace placeholders in text inside inline shapes."""
    for shape in doc.inline_shapes:
        if shape.text:
            for key, value in parameters.items():
                if f"{{{key}}}" in shape.text:
                    shape.text = shape.text.replace(f"{{{key}}}", str(value))

def update_lists(doc, parameters: dict[str, str]) -> None:
    """Replace placeholders in lists."""
    for paragraph in doc.paragraphs:
        if paragraph.style.name in ["List Paragraph", "Bullet"]:
            for key, value in parameters.items():
                if f"{{{key}}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(f"{{{key}}}", str(value))

def update_hyperlinks(doc, parameters: dict[str, str]) -> None:
    """Replace placeholders in hyperlinks."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if "http" in run.text:
                for key, value in parameters.items():
                    if f"{{{key}}}" in run.text:
                        run.text = run.text.replace(f"{{{key}}}", str(value))

def update_word_template(template_path: Path, output_path: Path, parameters: dict[str, str]) -> None:
    """Replace placeholders in all parts of a Word template with parameter values."""
    doc = Document(template_path)

    # Update main document content
    update_paragraphs(doc.paragraphs, parameters)
    update_tables(doc.tables, parameters)

    # Update headers and footers
    update_headers_footers(doc, parameters)

    # Update inline shapes
    update_inline_shapes(doc, parameters)

    # Update lists and hyperlinks
    update_lists(doc, parameters)
    update_hyperlinks(doc, parameters)

    doc.save(output_path)

def make_file_read_only(file_path: Path) -> None:
    """Set file to read-only."""
    os.chmod(file_path, 0o444)

def main() -> None:
    """Main script function."""
    # Define file paths
    base_dir = Path(__file__).resolve().parent.parent
    template_path = base_dir / "template/template-2.docx"
    excel_path = base_dir / "data/parameters-2.xlsx"
    output_path = base_dir / "template/output-2.docx"
    
    # Process files
    parameters = read_parameters(excel_path)
    update_word_template(template_path, output_path, parameters)
    make_file_read_only(output_path)
    print(f"Document with updated content saved to {output_path}")

if __name__ == "__main__":
    main()

  